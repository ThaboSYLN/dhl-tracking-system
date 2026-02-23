"""
Export services for generating PDF and DOCX reports
Handles document generation with tracking information

CHANGES MADE:
1. generate_pdf: Added binID column to PDF table (Lines 67-87, 91-102)
2. generate_docx: Added binID column to DOCX table (Lines 155-185, 189-210)
3. Both detailed and simple views now include binID
4. _filter_records: Added pre-transit filtering - records with pre-transit status
   are excluded from reports entirely before generation
"""
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
from datetime import datetime
import os
from pathlib import Path
import logging

from app.utils.config import settings
from app.models.database import TrackingRecord

logger = logging.getLogger(__name__)


# Status values that should be excluded from reports.
# All comparisons are done case-insensitively.
PRE_TRANSIT_STATUSES = {
    "pre-transit",
    "pretransit",
    "pre_transit",
}


class ExportService:
    """Service for exporting tracking data to PDF and DOCX"""

    def __init__(self):
        self.export_dir = settings.EXPORT_DIR
        Path(self.export_dir).mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _is_pre_transit(self, record: TrackingRecord) -> bool:
        """
        Return True if the record's status represents a pre-transit state.

        Checks both ``status_code`` and ``status`` fields so that whichever
        field the DHL API populates is caught.  Comparison is
        case-insensitive and strips surrounding whitespace.
        """
        fields_to_check = [
            record.status_code or "",
            record.status or "",
        ]
        for value in fields_to_check:
            normalised = value.strip().lower().replace(" ", "-")
            if normalised in PRE_TRANSIT_STATUSES:
                return True
        return False

    def _filter_records(
        self, records: List[TrackingRecord]
    ) -> List[TrackingRecord]:
        """
        Remove any records whose tracking status is pre-transit.

        Logs a summary so operators can see how many records were dropped
        without having to dig through the full record list.

        Args:
            records: Raw list of TrackingRecord objects.

        Returns:
            Filtered list with pre-transit records removed.
        """
        filtered = [r for r in records if not self._is_pre_transit(r)]
        dropped = len(records) - len(filtered)

        if dropped:
            dropped_numbers = [
                r.tracking_number
                for r in records
                if self._is_pre_transit(r)
            ]
            logger.info(
                f"Report filter: removed {dropped} pre-transit record(s) "
                f"from report → {dropped_numbers}"
            )
        else:
            logger.debug("Report filter: no pre-transit records found — all records included.")

        return filtered

    def _get_last_event_date(self, record: TrackingRecord) -> str:
        """Extract the most recent event timestamp from tracking details"""
        try:
            if record.tracking_details and isinstance(record.tracking_details, dict):
                events = record.tracking_details.get("events", [])
                if events and len(events) > 0:
                    most_recent_event = events[0]
                    timestamp = most_recent_event.get("timestamp")
                    if timestamp:
                        return timestamp

            if record.last_checked:
                return record.last_checked.strftime("%Y-%m-%dT%H:%M:%S+00:00")

            return "N/A"

        except Exception as e:
            logger.error(f"Error extracting last event date: {str(e)}")
            return "N/A"

    def generate_filename(self, format: str) -> str:
        """Generate unique filename for export"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tracking_report_{timestamp}.{format}"
        return os.path.join(self.export_dir, filename)

    # ------------------------------------------------------------------
    # PDF generation
    # ------------------------------------------------------------------

    def generate_pdf(
        self, tracking_records: List[TrackingRecord], include_details: bool = True
    ) -> str:
        """
        Generate PDF report.

        Pre-transit records are silently excluded before the document is
        built.  A note showing the original vs. included record count is
        added to the report header so readers know filtering occurred.

        Args:
            tracking_records: List of TrackingRecord objects.
            include_details: Include detailed origin/destination columns.

        Returns:
            Path to the generated PDF file.
        """
        try:
            original_count = len(tracking_records)
            tracking_records = self._filter_records(tracking_records)
            filtered_count = len(tracking_records)

            filename = self.generate_filename("pdf")
            doc = SimpleDocTemplate(filename, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()

            # ----------------------------------------------------------
            # Title
            # ----------------------------------------------------------
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                textColor=colors.HexColor("#1a1a1a"),
                spaceAfter=30,
                alignment=1,
            )
            title = Paragraph("DHL Tracking Report", title_style)
            elements.append(title)

            # ----------------------------------------------------------
            # Generation info  (includes filter note when records dropped)
            # ----------------------------------------------------------
            info_style = styles["Normal"]
            info_text = (
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
                f"Records Included: {filtered_count}"
            )
            if original_count != filtered_count:
                dropped = original_count - filtered_count
                info_text += (
                    f"<br/>"
                    f"<font color='#c0392b'>"
                    f"Note: {dropped} pre-transit record(s) excluded from this report."
                    f"</font>"
                )

            info = Paragraph(info_text, info_style)
            elements.append(info)
            elements.append(Spacer(1, 0.3 * inch))

            # ----------------------------------------------------------
            # Table data
            # ----------------------------------------------------------
            if not tracking_records:
                # Edge case: all records were pre-transit
                elements.append(
                    Paragraph(
                        "No records to display — all entries were pre-transit.",
                        info_style,
                    )
                )
                doc.build(elements)
                logger.info(f"PDF generated (empty after filter): {filename}")
                return filename

            if include_details:
                data = [[
                    "Tracking #",
                    "Bin ID",
                    "Status Code",
                    "Origin",
                    "Destination",
                    "Last Event Date",
                ]]
                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    data.append([
                        record.tracking_number,
                        record.bin_id or "N/A",
                        record.status_code or "N/A",
                        record.origin or "N/A",
                        record.destination or "N/A",
                        last_event_date,
                    ])
            else:
                data = [[
                    "Tracking #",
                    "Bin ID",
                    "Status Code",
                    "Last Event Date",
                ]]
                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    data.append([
                        record.tracking_number,
                        record.bin_id or "N/A",
                        record.status_code or "N/A",
                        last_event_date,
                    ])

            # ----------------------------------------------------------
            # Build & style table
            # ----------------------------------------------------------
            table = Table(data)
            table.setStyle(
                TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                    ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                    ("TEXTCOLOR", (0, 1), (-1, -1), colors.black),
                    ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    (
                        "ROWBACKGROUNDS",
                        (0, 1),
                        (-1, -1),
                        [colors.white, colors.lightgrey],
                    ),
                ])
            )

            elements.append(table)
            doc.build(elements)
            logger.info(f"PDF generated: {filename}")
            return filename

        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def generate_docx(
        self, tracking_records: List[TrackingRecord], include_details: bool = True
    ) -> str:
        """
        Generate DOCX report.

        Pre-transit records are silently excluded before the document is
        built.  A note showing how many records were filtered is added
        directly below the header paragraph.

        Args:
            tracking_records: List of TrackingRecord objects.
            include_details: Include detailed origin/destination columns.

        Returns:
            Path to the generated DOCX file.
        """
        try:
            original_count = len(tracking_records)
            tracking_records = self._filter_records(tracking_records)
            filtered_count = len(tracking_records)

            filename = self.generate_filename("docx")
            doc = Document()

            # ----------------------------------------------------------
            # Title
            # ----------------------------------------------------------
            title = doc.add_heading("DHL Tracking Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # ----------------------------------------------------------
            # Generation info
            # ----------------------------------------------------------
            info_para = doc.add_paragraph()
            info_para.add_run(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            ).bold = True
            info_para.add_run(
                f"Records Included: {filtered_count}"
            ).bold = True

            if original_count != filtered_count:
                dropped = original_count - filtered_count
                doc.add_paragraph()
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(
                    f"Note: {dropped} pre-transit record(s) were excluded from this report."
                )
                note_run.bold = True
                note_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)  # red

            doc.add_paragraph()

            # ----------------------------------------------------------
            # Edge case: nothing left after filtering
            # ----------------------------------------------------------
            if not tracking_records:
                doc.add_paragraph(
                    "No records to display — all entries were pre-transit."
                )
                doc.save(filename)
                logger.info(f"DOCX generated (empty after filter): {filename}")
                return filename

            # ----------------------------------------------------------
            # Table
            # ----------------------------------------------------------
            if include_details:
                table = doc.add_table(rows=1, cols=6)
                table.style = "Light Grid Accent 1"

                header_cells = table.rows[0].cells
                headers = [
                    "Tracking #",
                    "Bin ID",
                    "Status Code",
                    "Origin",
                    "Destination",
                    "Last Event Date",
                ]

                for idx, header in enumerate(headers):
                    cell = header_cells[idx]
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)

                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    row_cells = table.add_row().cells
                    row_cells[0].text = record.tracking_number
                    row_cells[1].text = record.bin_id or "N/A"
                    row_cells[2].text = record.status_code or "N/A"
                    row_cells[3].text = record.origin or "N/A"
                    row_cells[4].text = record.destination or "N/A"
                    row_cells[5].text = last_event_date

            else:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Grid Accent 1"

                header_cells = table.rows[0].cells
                headers = [
                    "Tracking #",
                    "Bin ID",
                    "Status Code",
                    "Last Event Date",
                ]

                for idx, header in enumerate(headers):
                    cell = header_cells[idx]
                    cell.text = header
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.size = Pt(10)

                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    row_cells = table.add_row().cells
                    row_cells[0].text = record.tracking_number
                    row_cells[1].text = record.bin_id or "N/A"
                    row_cells[2].text = record.status_code or "N/A"
                    row_cells[3].text = last_event_date

            doc.save(filename)
            logger.info(f"DOCX generated: {filename}")
            return filename

        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise

    # ------------------------------------------------------------------
    # Cleanup
    # ------------------------------------------------------------------

    def cleanup_old_exports(self, days: int = 7):
        """Clean up export files older than specified days"""
        try:
            cutoff_time = datetime.now().timestamp() - (days * 24 * 60 * 60)
            for file in os.listdir(self.export_dir):
                file_path = os.path.join(self.export_dir, file)
                if os.path.isfile(file_path):
                    if os.path.getmtime(file_path) < cutoff_time:
                        os.remove(file_path)
                        logger.info(f"Cleaned up old export: {file}")
        except Exception as e:
            logger.error(f"Error cleaning up exports: {str(e)}")


# Create export service instance
export_service = ExportService()