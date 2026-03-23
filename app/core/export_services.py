"""
Export services for generating PDF and DOCX reports
Handles document generation with tracking information

CHANGES MADE:
1. generate_pdf: Added binID column to PDF table
2. generate_docx: Added binID column to DOCX table
3. Both detailed and simple views now include binID
4. _filter_records: Added pre-transit filtering - records with pre-transit status
   are excluded from reports entirely before generation
5. Empty report alert: If ALL records are filtered out, an alert email is sent
   to thabospenser@gmail.com via EmailService.send_empty_report_alert()
6. generate_pdf / generate_docx: Added Waybill Creation Date column (date_order_binned)
7. PDF switched to landscape A4 with explicit column widths so all columns fit properly
8. Table size increased and original beige/lightgrey row colours restored
9. _abbreviate_location: Added location abbreviation lookup to shorten long
   origin/destination strings so they fit cleanly in the table columns.
   To add more abbreviations just add to the LOCATION_ABBREVIATIONS dictionary.
"""
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from typing import List, Dict, Any
from datetime import datetime
import os
from pathlib import Path
import logging

from app.utils.config import settings
from app.models.database import TrackingRecord

logger = logging.getLogger(__name__)


# -----------------------------------------------------------------------
# Add any status you want excluded from reports to this set.
# Comparisons are case-insensitive, so only lowercase entries are needed.
# -----------------------------------------------------------------------
PRE_TRANSIT_STATUSES = {
    "pre-transit",
    "pretransit",
    "pre_transit",
}


# -----------------------------------------------------------------------
# Location abbreviations dictionary.
# Keys are the full strings returned by the DHL API (uppercase).
# Values are the short codes shown in the report.
# Add as many as you need — matching is case-insensitive.
# -----------------------------------------------------------------------
LOCATION_ABBREVIATIONS = {
    # South Africa
    "JOHANNESBURG - SOUTH AFRICA":      "JHB - RSA",
    "CAPE TOWN - SOUTH AFRICA":         "CPT - RSA",
    "DURBAN - SOUTH AFRICA":            "DBN - RSA",
    "PRETORIA - SOUTH AFRICA":          "PTA - RSA",
    "PORT ELIZABETH - SOUTH AFRICA":    "PLZ - RSA",
    # Ghana
    "ACCRA - GHANA":                    "ACC - GHA",
    "ACCRA REMOTE - GHANA":             "ACC-R - GHA",
    "KUMASI - GHANA":                   "KMS - GHA",
    "GHANA REMOTE - GHANA":             "GHA-R - GHA",
    # Zambia
    "LUSAKA - ZAMBIA":                  "LUN - ZMB",
    # Zimbabwe
    "HARARE - ZIMBABWE":                "HRE - ZWE",
    "BULAWAYO - ZIMBABWE":              "BUQ - ZWE",
    # Botswana
    "GABORONE - BOTSWANA":              "GBE - BWA",
    # Malawi
    "LILONGWE - MALAWI":                "LLW - MWI",
    "BLANTYRE - MALAWI":                "BLZ - MWI",
    # Nigeria
    "LAGOS - NIGERIA":                  "LOS - NGA",
    "ABUJA - NIGERIA":                  "ABV - NGA",
    # Kenya
    "NAIROBI - KENYA":                  "NBO - KEN",
    # Tanzania
    "DAR ES SALAAM - TANZANIA":         "DAR - TZA",
    # Uganda
    "KAMPALA - UGANDA":                 "KLA - UGA",
    # Ethiopia
    "ADDIS ABABA - ETHIOPIA":           "ADD - ETH",
    # Mozambique
    "MAPUTO - MOZAMBIQUE":              "MPM - MOZ",
    # Angola
    "LUANDA - ANGOLA":                  "LAD - AGO",
    # Namibia
    "WINDHOEK - NAMIBIA":               "WDH - NAM",
    # Rwanda
    "KIGALI - RWANDA":                  "KGL - RWA",
    # Senegal
    "DAKAR - SENEGAL":                  "DKR - SEN",
    # Ivory Coast
    "ABIDJAN - IVORY COAST":            "ABJ - CIV",
    # Egypt
    "CAIRO - EGYPT":                    "CAI - EGY",
    # Morocco
    "CASABLANCA - MOROCCO":             "CMN - MAR",
}


class ExportService:
    """Service for exporting tracking data to PDF and DOCX"""

    def __init__(self):
        self.export_dir = settings.EXPORT_DIR
        Path(self.export_dir).mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _abbreviate_location(self, location: str) -> str:
        """
        Shorten a location string using the LOCATION_ABBREVIATIONS lookup.

        Matching is case-insensitive so 'Johannesburg - South Africa'
        and 'JOHANNESBURG - SOUTH AFRICA' both match.

        If no abbreviation is found the original value is returned unchanged
        so nothing ever shows as blank.

        To add a new location just add it to LOCATION_ABBREVIATIONS above.
        """
        if not location or location == 'N/A':
            return 'N/A'
        return LOCATION_ABBREVIATIONS.get(location.strip().upper(), location)

    def _is_pre_transit(self, record: TrackingRecord) -> bool:
        """
        Return True if the record's status represents a pre-transit state.

        Checks both status_code and status fields so that whichever
        field the DHL API populates is caught. Comparison is
        case-insensitive and strips surrounding whitespace.
        """
        fields_to_check = [
            record.status_code or "",
            record.status or "",
        ]
        for value in fields_to_check:
            normalised = value.strip().lower().replace(" ", "-")
            if normalised in PRE_TRANSIT_STATUSES:
                return True
        return False

    def _filter_records(self, records: List[TrackingRecord]) -> List[TrackingRecord]:
        """
        Remove any records whose tracking status is pre-transit.

        Logs a summary so operators can see how many records were dropped
        without having to dig through the full record list.
        """
        filtered = [r for r in records if not self._is_pre_transit(r)]
        dropped = len(records) - len(filtered)

        if dropped:
            dropped_numbers = [
                r.tracking_number
                for r in records
                if self._is_pre_transit(r)
            ]
            logger.info(
                f"Report filter: removed {dropped} pre-transit record(s) "
                f"from report → {dropped_numbers}"
            )
        else:
            logger.debug("Report filter: no pre-transit records found — all records included.")

        return filtered

    def _get_last_event_date(self, record: TrackingRecord) -> str:
        """Extract the most recent event timestamp from tracking details"""
        try:
            if record.tracking_details and isinstance(record.tracking_details, dict):
                events = record.tracking_details.get('events', [])

                if events and len(events) > 0:
                    most_recent_event = events[0]
                    timestamp = most_recent_event.get('timestamp')

                    if timestamp:
                        return timestamp

            if record.last_checked:
                return record.last_checked.strftime('%Y-%m-%dT%H:%M:%S+00:00')

            return 'N/A'

        except Exception as e:
            logger.error(f"Error extracting last event date: {str(e)}")
            return 'N/A'

    def _get_date_order_binned(self, record: TrackingRecord) -> str:
        """Return the waybill creation date (dateOrderBinned) or N/A."""
        return record.date_order_binned or "N/A"

    def generate_filename(self, format: str) -> str:
        """Generate unique filename for export"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"tracking_report_{timestamp}.{format}"
        return os.path.join(self.export_dir, filename)

    def _send_empty_alert(self, filename: str):
        """Fire the empty-report alert email."""
        try:
            from app.utils.email_service import EmailService
            EmailService().send_empty_report_alert(source_file=filename)
        except Exception as e:
            logger.error(f"Failed to send empty-report alert: {e}")

    # ------------------------------------------------------------------
    # PDF generation
    # ------------------------------------------------------------------

    def generate_pdf(self, tracking_records: List[TrackingRecord], include_details: bool = True) -> str:
        """
        Generate PDF report in landscape A4 with fixed column widths.
        Origin and Destination values are abbreviated using LOCATION_ABBREVIATIONS.

        Args:
            tracking_records: List of TrackingRecord objects
            include_details: Include detailed information

        Returns:
            Path to generated PDF file
        """
        try:
            original_count = len(tracking_records)
            tracking_records = self._filter_records(tracking_records)
            filtered_count = len(tracking_records)

            filename = self.generate_filename('pdf')

            page_size = landscape(A4)
            doc = SimpleDocTemplate(
                filename,
                pagesize=page_size,
                leftMargin=1.5*cm,
                rightMargin=1.5*cm,
                topMargin=1.5*cm,
                bottomMargin=1.5*cm,
            )
            elements = []
            styles = getSampleStyleSheet()

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1
            )
            title = Paragraph("DHL Tracking Report", title_style)
            elements.append(title)

            # Generation info
            info_style = styles['Normal']
            info_text = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br/>"
            info_text += f"Records Included: {filtered_count}"
            if original_count != filtered_count:
                dropped = original_count - filtered_count
                info_text += (
                    f"<br/>"
                    f"<font color='#c0392b'>"
                    f"Note: {dropped} pre-transit record(s) excluded from this report."
                    f"</font>"
                )
            info = Paragraph(info_text, info_style)
            elements.append(info)
            elements.append(Spacer(1, 0.3*inch))

            # Edge case: all records were filtered out
            if not tracking_records:
                elements.append(
                    Paragraph(
                        "No records to display — all entries were pre-transit.",
                        info_style,
                    )
                )
                doc.build(elements)
                logger.info(f"PDF generated (empty after filter): {filename}")
                self._send_empty_alert(filename)
                return filename

            # ----------------------------------------------------------
            # Column widths in cm
            # ----------------------------------------------------------
            if include_details:
                col_widths = [
                    3.5*cm,   # Tracking #
                    4.0*cm,   # Bin ID
                    4.0*cm,   # Waybill Creation Date
                    2.6*cm,   # Status Code
                    4.2*cm,   # Origin
                    4.2*cm,   # Destination
                    4.2*cm,   # Last Event Date
                ]
                data = [[
                    'Tracking #',
                    'Bin ID',
                    'Waybill Creation Date',
                    'Status Code',
                    'Origin',
                    'Destination',
                    'Last Event Date',
                ]]
                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    data.append([
                        record.tracking_number,
                        record.bin_id or 'N/A',
                        self._get_date_order_binned(record),
                        record.status_code or 'N/A',
                        self._abbreviate_location(record.origin or 'N/A'),
                        record.destination or 'N/A',
                        last_event_date,
                    ])
            else:
                col_widths = [
                    4.5*cm,   # Tracking #
                    4.5*cm,   # Bin ID
                    5.0*cm,   # Waybill Creation Date
                    3.5*cm,   # Status Code
                    9.2*cm,   # Last Event Date
                ]
                data = [[
                    'Tracking #',
                    'Bin ID',
                    'Waybill Creation Date',
                    'Status Code',
                    'Last Event Date',
                ]]
                for record in tracking_records:
                    last_event_date = self._get_last_event_date(record)
                    data.append([
                        record.tracking_number,
                        record.bin_id or 'N/A',
                        self._get_date_order_binned(record),
                        record.status_code or 'N/A',
                        last_event_date,
                    ])

            # Create table with explicit column widths
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                # Header row — original blue
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),
                # Data rows — original beige background
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
                ('TOPPADDING', (0, 1), (-1, -1), 10),
                # Alignment
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                # Grid
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                # Original alternating row colours — white and lightgrey
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
            ]))

            elements.append(table)
            doc.build(elements)
            logger.info(f"PDF generated: {filename}")
            return filename

        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise

    # ------------------------------------------------------------------
    # DOCX generation
    # ------------------------------------------------------------------

    def generate_docx(self, tracking_records: List[TrackingRecord], include_details: bool = True) -> str:
        """
        Generate DOCX report in landscape orientation.
        Origin and Destination values are abbreviated using LOCATION_ABBREVIATIONS.

        Args:
            tracking_records: List of TrackingRecord objects
            include_details: Include detailed information

        Returns:
            Path to generated DOCX file
        """
        try:
            original_count = len(tracking_records)
            tracking_records = self._filter_records(tracking_records)
            filtered_count = len(tracking_records)

            filename = self.generate_filename('docx')
            doc = Document()

            # Set landscape orientation
            section = doc.sections[0]
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

            # Title
            title = doc.add_heading('DHL Tracking Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Generation info
            info_para = doc.add_paragraph()
            info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = True
            info_para.add_run(f"Records Included: {filtered_count}").bold = True

            if original_count != filtered_count:
                dropped = original_count - filtered_count
                doc.add_paragraph()
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(
                    f"Note: {dropped} pre-transit record(s) were excluded from this report."
                )
                note_run.bold = True
                note_run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

            doc.add_paragraph()

            # Edge case: all records were filtered out
            if not tracking_records:
                doc.add_paragraph(
                    "No records to display — all entries were pre-transit."
                )
                doc.save(filename)
                logger.info(f"DOCX generated (empty after filter): {filename}")
                self._send_empty_alert(filename)
                return filename

            # Build table
            if include_details:
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Light Grid Accent 1'
                headers = [
                    'Tracking #',
                    'Bin ID',
                    'Waybill Creation Date',
                    'Status Code',
                    'Origin',
                    'Destination',
                    'Last Event Date',
                ]
                col_widths_cm = [3.5, 4.0, 4.0, 2.6, 4.2, 4.2, 4.2]
            else:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Light Grid Accent 1'
                headers = [
                    'Tracking #',
                    'Bin ID',
                    'Waybill Creation Date',
                    'Status Code',
                    'Last Event Date',
                ]
                col_widths_cm = [4.0, 4.0, 4.5, 3.2, 9.0]

            # Header row
            header_cells = table.rows[0].cells
            for idx, header in enumerate(headers):
                cell = header_cells[idx]
                cell.text = header
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(10)

            for idx, width in enumerate(col_widths_cm):
                header_cells[idx].width = Cm(width)

            # Data rows
            for record in tracking_records:
                last_event_date = self._get_last_event_date(record)
                row_cells = table.add_row().cells

                row_cells[0].text = record.tracking_number
                row_cells[1].text = record.bin_id or 'N/A'
                row_cells[2].text = self._get_date_order_binned(record)
                row_cells[3].text = record.status_code or 'N/A'

                if include_details:
                    row_cells[4].text = self._abbreviate_location(record.origin or 'N/A')
                    row_cells[5].text = record.destination or 'N/A'
                    row_cells[6].text = last_event_date
                else:
                    row_cells[4].text = last_event_date

                for idx, width in enumerate(col_widths_cm):
                    cell = row_cells[idx]
                    cell.width = Cm(width)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)

            doc.save(filename)
            logger.info(f"DOCX generated: {filename}")
            return filename

        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise

    def cleanup_old_exports(self, days: int = 7):
        """Clean up export files older than specified days"""
        try:
            cutoff_time = datetime.now().timestamp() - (days * 24 * 60 * 60)

            for file in os.listdir(self.export_dir):
                file_path = os.path.join(self.export_dir, file)
                if os.path.isfile(file_path):
                    if os.path.getmtime(file_path) < cutoff_time:
                        os.remove(file_path)
                        logger.info(f"Cleaned up old export: {file}")
        except Exception as e:
            logger.error(f"Error cleaning up exports: {str(e)}")


# Create export service instance
export_service = ExportService()